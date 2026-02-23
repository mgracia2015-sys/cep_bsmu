import streamlit as st
import warnings
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
from datetime import datetime
import io

# Для Google API
from googleapiclient.discovery import build
import google.auth

warnings.filterwarnings("ignore")

st.set_page_config(page_title="Редактор наукових статей", layout="centered")

st.title("📝 Автоматичне форматування статті")

# ============================================================
# 1️⃣ ІНТЕРФЕЙС STREAMLIT (Замість Radio Buttons Colab)
# ============================================================

col1, col2 = st.columns(2)

with col1:
    language_choice = st.radio(
        "Оберіть мову:",
        options=[('Українська', 'uk'), ('English', 'en')],
        format_func=lambda x: x[0]
    )
    language = language_choice[1]

with col2:
    article_type_choice = st.radio(
        "Тип статті:",
        options=[
            ('Оригінальне дослідження', 'original'),
            ('Клінічний випадок', 'case'),
            ('Огляд літератури', 'review')
        ],
        format_func=lambda x: x[0]
    )
    article_type = article_type_choice[1]

uploaded_file = st.file_uploader("Завантажте файл .docx", type=["docx"])

# Кнопка запуску
if st.button("🚀 Обробити статтю") and uploaded_file is not None:
    
    # ============================================================
    # 2️⃣ ГОЛОВНА ЛОГІКА (БЕЗ ЗМІН В АЛГОРИТМІ)
    # ============================================================
    
    report = []
    
    # Завантаження
    doc = Document(uploaded_file)
    file_name = uploaded_file.name
    report.append("Файл завантажено: " + file_name)

    # 2.3 Перевірка та виправлення полів сторінки
    section = doc.sections[0]
    if section.top_margin != Cm(2):
        section.top_margin = Cm(2)
        report.append("Виправлено верхнє поле на 2 см")
    if section.bottom_margin != Cm(2):
        section.bottom_margin = Cm(2)
        report.append("Виправлено нижнє поле на 2 см")
    if section.right_margin != Cm(2):
        section.right_margin = Cm(2)
        report.append("Виправлено праве поле на 2 см")
    if section.left_margin != Cm(2):
        section.left_margin = Cm(2)
        report.append("Виправлено ліве поле на 2 см")

    # 2.4 Перевірка та виправлення формату тексту
    for paragraph in doc.paragraphs:
        paragraph_format = paragraph.paragraph_format
        paragraph_format.line_spacing = 1.5
        paragraph_format.space_before = Pt(0)
        paragraph_format.space_after = Pt(0)
        if paragraph.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY:
            paragraph_format.first_line_indent = Cm(1.25)
        for run in paragraph.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(14)
    report.append("Перевірено та виправлено формат тексту")

    # 2.5 УДК, назва та автори
    paragraphs = doc.paragraphs

    # ---- УДК ----
    if len(paragraphs) >= 1:
        first = paragraphs[0]
        if not first.text.startswith("УДК"):
            first.text = "УДК 000.00"
            report.append("Додано УДК")
        for run in first.runs:
            run.font.bold = True
        report.append("УДК перевірено/виправлено")

    # ---- Назва статті ----
    if len(paragraphs) >= 2:
        title_para = paragraphs[1]
        title_text = title_para.text.replace("\n", " ").strip().upper()
        title_para.text = title_text
        for run in title_para.runs:
            run.font.bold = True
        report.append("Назва статті перевірена та приведена до формату (великими літерами, один абзац, жирний)")

        # ---- Автори ----
        if len(paragraphs) >= 3:
            authors_para = paragraphs[2]
            authors_list = authors_para.text.split(',')
            new_authors = []
            for author in authors_list:
                author = author.strip()
                parts = author.split()
                if len(parts) >= 2:
                    if parts[0].endswith("."): 
                        initials = parts[0]
                        surname = parts[1]
                        rest = " ".join(parts[2:])
                    else:
                        surname = parts[0]
                        initials = parts[1]
                        rest = " ".join(parts[2:])
                    author_text = f"{initials} {surname}"
                    if rest: author_text += f" {rest}"
                    new_authors.append(author_text)
                else:
                    new_authors.append(author)
            
            authors_para.text = ", ".join(new_authors)
            for run in authors_para.runs:
                run.font.bold = True
                run.font.italic = True
            report.append("Автори перевірені та відформатовані (жирний + курсив, ініціали перед прізвищем, цифри афіліацій збережені)")

    # Афіліація та анотація
    max_affiliation_number = 0
    if len(paragraphs) >= 3:
        authors_para = paragraphs[2]
        numbers = re.findall(r'\d+', authors_para.text)
        if numbers:
            max_affiliation_number = max([int(n) for n in numbers])

    affiliation_start = 3
    affiliation_end = affiliation_start + max_affiliation_number if max_affiliation_number > 0 else affiliation_start + 1
    
    affiliation_paragraphs = paragraphs[affiliation_start:affiliation_end]
    for para in affiliation_paragraphs:
        for run in para.runs:
            run.font.bold = False
            run.font.italic = False
    report.append(f"Афіліація авторів перевірена ({affiliation_end - affiliation_start} рядків)")

    # Анотація
    abstract_start = affiliation_end
    abstract_end = abstract_start
    keywords_found = False
    for i in range(abstract_start, len(paragraphs)):
        para_text_lower = paragraphs[i].text.lower()
        if "ключові слова" in para_text_lower or "keywords" in para_text_lower:
            abstract_end = i + 1
            keywords_found = True
            break
    if not keywords_found: abstract_end = len(paragraphs)

    abstract_paragraphs = paragraphs[abstract_start:abstract_end]
    abstract_text = ""
    for para in abstract_paragraphs:
        abstract_text += para.text + "\n"
        para.paragraph_format.first_line_indent = None
        for run in para.runs:
            current_bold = run.font.bold
            run.font.italic = True
            if current_bold is not None: run.font.bold = current_bold

    abstract_length = len(abstract_text)
    if abstract_length < 1800 or abstract_length > 2500:
        report.append(f"⚠️ Попередження: довжина анотації {abstract_length} символів (рекомендовано 1800–2500)")
    report.append("Анотація перевірена та відформатована (курсив)")

    # 2.X ДРУГА МОВНА ВЕРСІЯ
    second_start = abstract_end
    while second_start < len(paragraphs) and not paragraphs[second_start].text.strip():
        second_start += 1

    if second_start < len(paragraphs):
        title2_para = paragraphs[second_start]
        title2_text = title2_para.text.replace("\n", " ").strip().upper()
        title2_para.text = title2_text
        for run in title2_para.runs: run.font.bold = True
        report.append("Назва другою мовою перевірена та приведена до формату")

        authors2_index = second_start + 1
        while authors2_index < len(paragraphs) and not paragraphs[authors2_index].text.strip():
            authors2_index += 1

        if authors2_index < len(paragraphs):
            authors2_para = paragraphs[authors2_index]
            authors_list = authors2_para.text.split(',')
            new_authors = []
            for author in authors_list:
                author = author.strip()
                parts = author.split()
                if len(parts) >= 2:
                    if parts[0].endswith("."):
                        initials, surname = parts[0], parts[1]
                        rest = " ".join(parts[2:])
                    else:
                        surname, initials = parts[0], parts[1]
                        rest = " ".join(parts[2:])
                    author_text = f"{initials} {surname}"
                    if rest: author_text += f" {rest}"
                    new_authors.append(author_text)
                else: new_authors.append(author)
            authors2_para.text = ", ".join(new_authors)
            for run in authors2_para.runs:
                run.font.bold, run.font.italic = True, True
            report.append("Автори другою мовою відформатовані")

            affiliation2_start = authors2_index + 1
            affiliation2_end = affiliation2_start + (affiliation_end - affiliation_start)
            affiliation2_paragraphs = paragraphs[affiliation2_start:affiliation2_end]
            for para in affiliation2_paragraphs:
                for run in para.runs:
                    run.font.bold, run.font.italic = False, False
            report.append(f"Афіліація другою мовою перевірена ({len(affiliation2_paragraphs)} рядків)")

            abstract2_start = affiliation2_end
            abstract2_end = abstract2_start
            keywords2_found = False
            for i in range(abstract2_start, len(paragraphs)):
                para_text_lower = paragraphs[i].text.lower()
                if "ключові слова" in para_text_lower or "keywords" in para_text_lower:
                    abstract2_end = i + 1
                    keywords2_found = True
                    break
            if not keywords2_found: abstract2_end = len(paragraphs)
            
            abstract2_paragraphs = paragraphs[abstract2_start:abstract2_end]
            for para in abstract2_paragraphs:
                para.paragraph_format.first_line_indent = None
                for run in para.runs:
                    current_bold = run.font.bold
                    run.font.italic = True
                    if current_bold is not None: run.font.bold = current_bold
            report.append("Анотація другою мовою перевірена та відформатована (курсив)")

    # 3. Перевірка наявності структурних елементів
    required_elements = {
        "uk": { 
            "original": {
                "abstract_uk": ["Мета дослідження", "Матеріали і методи", "Результати", "Висновки", "Ключові слова"],
                "abstract_en": ["Objective", "Materials and methods", "Results", "Conclusions", "Key words"],
                "main_text": ["Вступ", "Мета роботи", "Матеріал і методи дослідження", "Результати та їх обговорення", "Висновки", "Перспективи подальших досліджень", "Список літератури", "References"]
            },
            "case": {
                "abstract_uk": ["Висновки"], "abstract_en": ["Conclusions"],
                "main_text": ["Вступ", "Опис клінічного випадку", "Висновки", "Список літератури", "References"]
            },
            "review": {
                "abstract_uk": ["Мета роботи", "Основна частина", "Висновки"],
                "abstract_en": ["Objective", "Main Text", "Conclusions"],
                "main_text": ["Вступ", "Мета роботи", "Основна частина", "Висновки", "Список літератури", "References"]
            }
        },
        "en": {
            "original": {
                "abstract_en": ["Objective", "Materials and methods", "Results", "Conclusions", "Key words"],
                "abstract_uk": ["Мета дослідження", "Матеріали і методи", "Результати", "Висновки", "Ключові слова"],
                "main_text": ["Introduction", "Objective", "Materials and Methods", "Results and Discussion", "Conclusions", "Prospects for further research", "References"]
            },
            "case": {
                "abstract_en": ["Conclusions"], "abstract_uk": ["Висновки"],
                "main_text": ["Introduction", "Case description", "Conclusions", "References"]
            },
            "review": {
                "abstract_en": ["Objective", "Materials and methods", "Results", "Conclusions", "Key words"],
                "abstract_uk": ["Мета дослідження", "Матеріали і методи", "Результати", "Висновки", "Ключові слова"],
                "main_text": ["Introduction", "Objective", "Main part", "Conclusions", "References"]
            }
        }
    }

    selected_structure = required_elements[language][article_type]
    missing_elements = []
    for section_name, elements in selected_structure.items():
        for element in elements:
            found = False
            for paragraph in doc.paragraphs:
                if paragraph.text.strip().lower().startswith(element.lower()):
                    found = True
                    break
            if not found: missing_elements.append(f"{section_name}: {element}")

    if missing_elements:
        report.append("❌ Відсутні або неправильно оформлені структурні елементи:")
        for item in missing_elements: report.append(f"   - {item}")
    else: report.append("✅ Усі обов’язкові структурні елементи присутні та оформлені правильно")

    # ПЕРЕВІРКА ЛІТЕРАТУРИ
    references_start, references_title = None, None
    titles_to_find = ["список літератури", "references"]
    for i, para in enumerate(paragraphs):
        text_lower = para.text.strip().lower()
        if any(text_lower.startswith(t) for t in titles_to_find):
            references_start, references_title = i + 1, para.text.strip()
            break

    if references_start is None:
        report.append("❌ Не знайдено розділ літератури")
    else:
        reference_paragraphs = []
        for para in paragraphs[references_start:]:
            text = para.text.strip()
            if not text: continue
            if re.search(r"(author|email|e-mail|correspondence|адреса|контакт)", text.lower()): break
            reference_paragraphs.append(text)
        
        reference_count = len(reference_paragraphs)
        if article_type in ["original", "case"]:
            if reference_count > 15: report.append(f"⚠️ Джерел: {reference_count} (допустимо не більше 15)")
            else: report.append(f"Кількість джерел: {reference_count}")
        elif article_type == "review":
            if reference_count < 50: report.append(f"⚠️ Джерел: {reference_count} (для огляду потрібно не менше 50)")
            else: report.append(f"Кількість джерел: {reference_count}")

        expected_number, numbering_errors, vancouver_errors = 1, False, False
        for ref in reference_paragraphs:
            match = re.match(r"^(\d+)[\.\)]", ref)
            if match:
                if int(match.group(1)) != expected_number: numbering_errors = True
                expected_number += 1
            else: numbering_errors = True
            if not re.search(r"\b(19|20)\d{2}\b", ref): vancouver_errors = True
        
        if vancouver_errors: report.append("⚠️ Можливе порушення Vancouver style")
        else: report.append("Стиль літератури виглядає коректним (базова перевірка)")

    # 2.6 Збереження файлу в пам'ять для завантаження
    bio = io.BytesIO()
    doc.save(bio)
    
    # 2.7 Відображення звіту в Streamlit
    st.subheader("=== ЗВІТ ПРО ВНЕСЕНІ ЗМІНИ ===")
    
    sections = {"Файл": [], "Поля сторінки": [], "Формат тексту": [], "Назва/УДК/Автори": [], "Інше": []}
    for item in report:
        if "Файл завантажено" in item: sections["Файл"].append(item)
        elif "поле" in item: sections["Поля сторінки"].append(item)
        elif any(x in item for x in ["шрифт", "міжрядковий", "відступ", "формат тексту"]): sections["Формат тексту"].append(item)
        elif any(x in item for x in ["УДК", "Назва", "Автори"]): sections["Назва/УДК/Автори"].append(item)
        else: sections["Інше"].append(item)

    for sec, items in sections.items():
        if items:
            with st.expander(f"📌 {sec}", expanded=True):
                for it in list(dict.fromkeys(items)):
                    st.write(f"- {it}")

    st.success("Готово ✅ Файл відформатовано.")
    
    # Кнопка завантаження
    st.download_button(
        label="📥 Завантажити виправлений файл",
        data=bio.getvalue(),
        file_name=f"fixed_{file_name}",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

    # 2.8 Google Docs Log (Спроба виконати, якщо є сертифікати)
    # ПРИМІТКА: У Streamlit Cloud авторизація через auth.authenticate_user() не працюватиме як у Colab.
    # Потрібен файл service_account.json або налаштовані Secrets.
    try:
        # Ця частина залишиться робочою ТІЛЬКИ якщо запущено локально з наявним браузером
        # або налаштованими змінними оточення Google.
        # В Streamlit Cloud вона швидше за все видасть помилку без додаткових налаштувань.
        # auth.authenticate_user()  <-- Видалено, бо це специфічно для Colab
        creds, _ = google.auth.default()
        service = build('docs', 'v1', credentials=creds)
        current_date = datetime.now().strftime("%d.%m.%Y")
        udk_title = paragraphs[1].text if len(paragraphs) > 0 else "Невідомо"
        authors = paragraphs[2].text if len(paragraphs) > 2 else "Невідомо"
        text_to_insert = f"\n[{current_date}] АВТОР: {authors} | СТАТТЯ: {udk_title}\n"
        requests = [{'insertText': {'location': {'index': 1}, 'text': text_to_insert}}]
        SHARED_LOG_DOC_ID = '13j6RQGukjUHqTu4doCFeqVtS7PlbrfBIKXVG8Kg7qzo'
        service.documents().batchUpdate(documentId=SHARED_LOG_DOC_ID, body={'requests': requests}).execute()
        st.info("✅ Дані успішно додані у Google Docs")
    except Exception as e:
        st.warning(f"⚠️ Запис в журнал Google Docs пропущено (потрібне налаштування доступу): {e}")

elif uploaded_file is None:
    st.info("Будь ласка, завантажте файл, щоб почати.")